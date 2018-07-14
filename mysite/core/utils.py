from docx import Document
from docx.shared import Inches
import os
from .models import Docs

def saveToDoc(id):
    # doc=Document()
    # doc.add_heading("TEST",0)
    # doc.add_paragraph("test Docs")
    # doc.add_page_break()

    # path = 'documents/user_{0}/{1}.docx'.format(instance.customer.id, instance.doc_type)
    # # doc.save(path)
    #
    # try:
    #     doc = Document(path)
    # except Exception as e:
    #     doc = Document()
    #     print(e)
    #
    # doc.add_picture('documents/user_8/Screenshot_from_2018-04-15_031105.png')
    # doc.save(path)

    import pdb
    pdb.set_trace()




    docs=Docs.objects.filter(customer__id=id)
    doc_path='documents/user_{0}/all_docs.docx'.format(id)
    doc_file=Document()

    from os import listdir
    from os.path import isfile, join
    mypath='documents/user_{0}/'.format(id)
    onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]

    for f in onlyfiles:
        ext=os.path.splitext(f)[1]
        valid_extensions = ['.png', '.jpg', '.jpeg']
        if  ext.lower() in valid_extensions:
            path="documents/user_{0}/{1}".format(id,f)
            doc_file.add_picture(path,width=Inches(4.0))
            doc_file.add_page_break()


    # for doc in docs:
    #     path = os.path.abspath(doc.doc_file.name)
    #     doc_file.add_picture(path,width=Inches(4.0))
    #     doc_file.add_page_break()

    doc_file.save(doc_path)



    return

