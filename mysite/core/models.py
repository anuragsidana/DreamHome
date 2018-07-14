from __future__ import unicode_literals

from django.db import models
from django.contrib.auth.models import User
from django.db.models.signals import post_save
from django.dispatch import receiver
from django.urls import reverse




class Profile(models.Model):
    user = models.OneToOneField(User, on_delete=models.CASCADE)
    bio = models.TextField(max_length=500, blank=True)
    location = models.CharField(max_length=30, blank=True)
    birth_date = models.DateField(null=True, blank=True)
    is_admin = models.BooleanField(default=False)
    # ADMIN_CHOICE = (("YES", True), ("NO", False))
    # is_admin = models.BooleanField(default=False)

    def __str__(self):
        return self.user.username

#signal when user created then also create profile object
@receiver(post_save, sender=User)
def update_user_profile(sender, instance, created, **kwargs):
    if created:
        Profile.objects.create(user=instance)
    instance.profile.save()


class Customer(models.Model):
    agent=models.ForeignKey(Profile,on_delete=models.SET_NULL,null=True)
    first_name=models.CharField(max_length=200)
    last_name=models.CharField(max_length=200,blank=True)
    # father_name=models.CharField(max_length=200)
    # phone_regex = RegexValidator(regex=r'^\+?1?\d{9,15}$',
    #                              message="Phone number must be entered in the format: '+999999999'. Up to 15 digits allowed.")
    # phone_number = models.CharField(validators=[phone_regex], max_length=17, blank=True)  # validators should be a list
    # permanent_address=models.CharField(max_length=500)
    # aadhar_regx= RegexValidator(regex= r'^\d{12}$',message='aadhar must be of 121 digits')
    # aadhar_no= models.CharField(validators=[aadhar_regx],max_length=12)
    # # whenever a new Cutomer created redirect it to its details page
    def get_absolute_url(self):
        # detail view need primary key argument , hence to get primary key of the self object we wrote kwargs
        return reverse('customer_detail', kwargs={'pk': self.pk})

    def __str__(self):
        return self.first_name

import os
#https://stackoverflow.com/questions/2680391/in-django-changing-the-file-name-of-an-uploaded-file
# def update_filename(name):
#
#     def wrapper(instance,filename):
#         path = "documents/%Y/%m/%d/" + instance.doc_type
#         format = instance.userid + instance.transaction_uuid + instance.file_extension
#         return os.path.join(path, name)
#
#     return wrapper
#
#
# import os
# from uuid import uuid4
#
# def path_and_rename(path):
#     def wrapper(instance, filename):
#         ext = filename.split('.')[-1]
#         # get filename
#         if instance.pk:
#             filename = '{}.{}'.format(instance.pk, ext)
#         else:
#             # set filename as random string
#             filename = '{}.{}'.format(uuid4().hex, ext)
#         # return the whole path to the file
#         return os.path.join(path, filename)
#     return wrapper


from docx import Document
from docx.shared import Inches
def user_directory_path(instance, filename):
    path = 'documents/user_{0}/{1}.docx'.format(instance.customer.id, instance.doc_type)
    # doc.save(path)

    try:
        doc = Document(path)
        doc.add_page_break()
    except Exception as e:
        doc = Document()
        print(e)
    #
    # doc.add_picture('documents/user_8/Screenshot_from_2018-04-15_031105.png',width=Inches(4.0))
    # doc.save(path)

    # file will be uploaded to MEDIA_ROOT/user_<id>/<filename>
    return 'documents/user_{0}/{1}'.format(instance.customer.id, filename)



class Docs(models.Model):
    customer=models.ForeignKey(Customer,on_delete=models.SET_NULL,null=True)
    doc_type=models.CharField(max_length=200)
    doc_file=models.FileField(null=True,upload_to=user_directory_path)


    # whenever a new Document created redirect it to its details page
    def get_absolute_url(self):
        # detail view need primary key argument , hence to get primary key of the self object we wrote kwargs
        return reverse('document_detail', kwargs={'pk': self.pk})

    def __str__(self):
        return self.doc_type



